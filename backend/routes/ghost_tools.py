"""
Ghost AI Tools — file creation, web intelligence, code interpreter, data analysis.

Routes:
  POST /api/ghost/tools/create-file         — create PDF/XLSX/DOCX/CSV/TXT
  POST /api/ghost/tools/create-presentation — create slide deck
  POST /api/ghost/tools/web-search          — search web, return results
  POST /api/ghost/tools/read-url            — scrape and summarize a URL
  POST /api/ghost/tools/read-pdf            — read uploaded PDF, summarize
  POST /api/ghost/tools/stock-report        — stock analysis
  POST /api/ghost/tools/weather-report      — weather forecast
  POST /api/ghost/tools/news                — latest news on topic
  POST /api/ghost/tools/translate-doc       — translate document
  POST /api/ghost/tools/run-code            — execute Python code
  POST /api/ghost/tools/analyze-data        — analyze CSV data
"""

from __future__ import annotations
import os
import io
import re
import json
import base64
import tempfile
import uuid
import logging
import httpx
from datetime import datetime
from fastapi import APIRouter, HTTPException, Header, UploadFile, File, Form, Depends
from fastapi.responses import JSONResponse
from pydantic import BaseModel
from typing import Optional
from sqlalchemy.ext.asyncio import AsyncSession
from db import get_db

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/ghost/tools", tags=["ghost-tools"])

ANTHROPIC_KEY = os.getenv("ANTHROPIC_API_KEY", "")

# Lazy import to avoid circular dependency
def _audit_log_lazy():
    from routes.ghost_auth import _audit_log
    return _audit_log

# In-memory file cache — backed by R2 for persistence across restarts
FILE_STORE: dict[str, dict] = {}
_FILE_STORE_MAX = 500  # max files in memory


async def _store_file(file_id: str, filename: str, mime: str, file_bytes: bytes, owner_email: str = ""):
    """Store file in memory cache AND upload to R2 for persistence."""
    # Evict oldest files if cache is full
    if len(FILE_STORE) >= _FILE_STORE_MAX:
        oldest = min(FILE_STORE, key=lambda k: FILE_STORE[k].get("created", ""))
        del FILE_STORE[oldest]
    file_b64 = base64.b64encode(file_bytes).decode()
    FILE_STORE[file_id] = {"filename": filename, "mime": mime, "data": file_b64, "created": datetime.utcnow().isoformat(), "owner_email": owner_email}
    # Persist to R2 (best-effort, non-fatal)
    try:
        from utils.file_storage import upload_to_r2
        await upload_to_r2(f"ghost-files/{file_id}", file_bytes, content_type=mime, filename=filename)
    except Exception as e:
        import logging
        logging.getLogger(__name__).warning(f"R2 upload failed for {file_id}: {e}")


async def _get_file(file_id: str) -> dict | None:
    """Get file from in-memory cache, falling back to R2."""
    if file_id in FILE_STORE:
        return FILE_STORE[file_id]
    # Try R2
    try:
        from utils.file_storage import download_from_r2
        result = await download_from_r2(f"ghost-files/{file_id}")
        if result:
            file_b64 = base64.b64encode(result["data"]).decode()
            entry = {"filename": result["filename"] or f"file_{file_id[:8]}", "mime": result["content_type"] or "application/octet-stream", "data": file_b64, "created": datetime.utcnow().isoformat()}
            FILE_STORE[file_id] = entry  # Re-populate cache
            return entry
    except Exception:
        pass
    return None


def _verify_auth(authorization: str):
    from routes.ghost_auth import verify_ghost_token
    token = authorization.replace("Bearer ", "") if authorization.startswith("Bearer ") else authorization
    return verify_ghost_token(token)


async def _ask_claude(prompt: str, system: str = "You are a helpful assistant.") -> str:
    if not ANTHROPIC_KEY:
        raise HTTPException(500, "API key not configured")
    import time as _time
    for attempt in range(2):
        t0 = _time.time()
        logger.info("[_ask_claude] attempt=%d prompt_len=%d system_len=%d", attempt, len(prompt), len(system))
        try:
            async with httpx.AsyncClient(timeout=120) as client:
                res = await client.post(
                    "https://api.anthropic.com/v1/messages",
                    headers={"Content-Type": "application/json", "x-api-key": ANTHROPIC_KEY, "anthropic-version": "2023-06-01"},
                    json={"model": "claude-sonnet-4-20250514", "max_tokens": 4096, "system": system, "messages": [{"role": "user", "content": prompt}]},
                )
                elapsed = _time.time() - t0
                logger.info("[_ask_claude] attempt=%d status=%d elapsed=%.1fs", attempt, res.status_code, elapsed)
                if res.status_code == 529:
                    logger.warning("[_ask_claude] API overloaded (529), retrying in 3s...")
                    if attempt == 0:
                        import asyncio
                        await asyncio.sleep(3)
                        continue
                if res.status_code != 200:
                    body = res.text[:500]
                    logger.error("[_ask_claude] API error %d: %s", res.status_code, body)
                    raise HTTPException(res.status_code, f"AI error: {res.status_code}")
                result = res.json().get("content", [{}])[0].get("text", "")
                logger.info("[_ask_claude] success, response_len=%d elapsed=%.1fs", len(result), elapsed)
                return result
        except httpx.TimeoutException:
            elapsed = _time.time() - t0
            logger.error("[_ask_claude] TIMEOUT after %.1fs attempt=%d", elapsed, attempt)
            if attempt == 0:
                continue
            raise HTTPException(504, "AI request timed out after 120s")
        except HTTPException:
            raise
        except Exception as e:
            elapsed = _time.time() - t0
            logger.error("[_ask_claude] unexpected error after %.1fs: %s", elapsed, str(e))
            raise
    raise HTTPException(500, "AI request failed after retries")


# ─── FILE CREATION ────────────────────────────────────────────────────────

class CreateFileRequest(BaseModel):
    description: str  # What the file should contain
    file_type: Optional[str] = "pdf"  # pdf, xlsx, docx, csv, txt
    filename: Optional[str] = None
    quality: Optional[str] = "standard"  # standard or premium


@router.post("/create-file")
async def create_file(req: CreateFileRequest, authorization: str = Header(...), db: AsyncSession = Depends(get_db)):
    payload = _verify_auth(authorization)

    # ── Check if this is a structured document type (invoice, resume, etc.) ──
    if req.file_type == "pdf":
        from lib.pdf_document_templates import detect_document_type, get_structured_prompt, render_structured_pdf
        doc_type = detect_document_type(req.description)
        if doc_type:
            structured_prompt = get_structured_prompt(doc_type)
            raw_json = await _ask_claude(req.description, structured_prompt)
            # Strip markdown fences if Claude wraps in ```json
            clean = raw_json.strip()
            if clean.startswith('```'):
                clean = clean.split('\n', 1)[1] if '\n' in clean else clean[3:]
                if clean.endswith('```'):
                    clean = clean[:-3]
                clean = clean.strip()
            try:
                data = json.loads(clean)
                file_bytes = render_structured_pdf(doc_type, data)
                file_id = str(uuid.uuid4())
                filename = req.filename or f"{doc_type}_{file_id[:8]}"
                filename += ".pdf"
                mime = "application/pdf"
                await _store_file(file_id, filename, mime, file_bytes)
                await _audit_log_lazy()(db, payload.get("email", ""), "tool_create_file", f"Structured {doc_type} PDF")
                await db.commit()
                return {"file_id": str(file_id), "filename": filename, "download_url": f"/api/ghost/tools/download/{file_id}"}
            except (json.JSONDecodeError, Exception) as e:
                logger.warning("[create-file] Structured %s failed (%s), falling back to generic", doc_type, e)

    # Ask AI to generate the content with professional quality
    format_instructions = {
        "csv": "Return raw CSV data with headers. Use proper column names. Include realistic, detailed data.",
        "txt": "Return well-written plain text with clear structure.",
        "xlsx": "Return a JSON array of objects where keys are column headers. Include realistic data with at least 10 rows. Use proper number formatting.",
        "pdf": "Return well-structured text using markdown-style headings (# for H1, ## for H2, ### for H3). Use - for bullet points. Use **bold** for emphasis. Use | column1 | column2 | format for tables. Use --- for section dividers. Include specific numbers, metrics, and details.",
        "docx": "Return well-structured text using markdown-style headings (# for H1, ## for H2, ### for H3). Use - for bullet points.",
    }

    system = f"""You are an expert professional writer and document creator. You write with the quality of a top-tier consultant at McKinsey, Goldman Sachs, or a Big 4 firm.

DOCUMENT TYPE: {req.file_type.upper()}

WRITING RULES:
- Write with authority, precision, and professionalism
- Use industry-standard terminology and frameworks
- Include specific details, numbers, metrics, and examples — never vague
- Structure content with clear hierarchy: sections, subsections, bullet points
- For resumes: use strong action verbs, quantify achievements, tailor to the role
- For business documents: include executive summary, key findings, recommendations
- For proposals: lead with value proposition, include timeline, budget, deliverables
- For reports: use data-driven language, cite sources where relevant
- For invoices: use proper formatting with line items, rates, subtotals, tax, total
- Content should be thorough and detailed — aim for 2-4 pages of real content
- Write as if this document will be presented to a CEO or client

FORMAT: {format_instructions.get(req.file_type, 'Return well-structured text with headings.')}

Return ONLY the document content. No explanations, no preamble, no "here is your document". Start directly with the content."""

    content = await _ask_claude(req.description, system)

    file_id = str(uuid.uuid4())
    filename = req.filename or f"document_{file_id[:8]}"

    if req.file_type == "csv":
        file_bytes = content.encode('utf-8')
        filename += ".csv"
        mime = "text/csv"

    elif req.file_type == "txt":
        file_bytes = content.encode('utf-8')
        filename += ".txt"
        mime = "text/plain"

    elif req.file_type == "xlsx":
        try:
            import openpyxl
            wb = openpyxl.Workbook()
            ws = wb.active

            # Try to parse JSON data from AI
            try:
                data = json.loads(content)
                if isinstance(data, list) and len(data) > 0:
                    headers = list(data[0].keys())
                    ws.append(headers)
                    for row in data:
                        ws.append([row.get(h, "") for h in headers])
            except json.JSONDecodeError:
                # Fallback: put content as text
                for i, line in enumerate(content.split('\n')):
                    cells = line.split(',') if ',' in line else [line]
                    ws.append(cells)

            buf = io.BytesIO()
            wb.save(buf)
            file_bytes = buf.getvalue()
            filename += ".xlsx"
            mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        except ImportError:
            # Fallback to CSV if openpyxl not installed
            file_bytes = content.encode('utf-8')
            filename += ".csv"
            mime = "text/csv"

    elif req.file_type == "docx":
        try:
            from docx import Document
            doc = Document()
            for line in content.split('\n'):
                line = line.strip()
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('- '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif line:
                    doc.add_paragraph(line)
            buf = io.BytesIO()
            doc.save(buf)
            file_bytes = buf.getvalue()
            filename += ".docx"
            mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        except ImportError:
            file_bytes = content.encode('utf-8')
            filename += ".txt"
            mime = "text/plain"

    else:  # pdf
        try:
            from lib.pdf_templates import create_professional_pdf
            file_bytes = create_professional_pdf(content, title=req.filename)
            filename += ".pdf"
            mime = "application/pdf"
        except Exception as e:
            logger.error("PDF creation failed: %s", e)
            # Fallback: still try basic PDF with reportlab
            try:
                from reportlab.lib.pagesizes import letter
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                from reportlab.lib.styles import getSampleStyleSheet
                buf = io.BytesIO()
                doc = SimpleDocTemplate(buf, pagesize=letter)
                styles = getSampleStyleSheet()
                story = []
                for line in content.split('\n'):
                    line = line.strip()
                    if not line:
                        story.append(Spacer(1, 12))
                    else:
                        # Strip markdown bold/italic
                        clean = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
                        clean = re.sub(r'\*(.+?)\*', r'\1', clean)
                        clean = re.sub(r'^#+\s*', '', clean)
                        story.append(Paragraph(clean, styles['Normal']))
                doc.build(story)
                file_bytes = buf.getvalue()
                filename += ".pdf"
                mime = "application/pdf"
            except Exception as e2:
                logger.error("PDF fallback also failed: %s", e2)
                file_bytes = content.encode('utf-8')
                filename += ".txt"
                mime = "text/plain"

    # Store file (memory + R2)
    await _store_file(file_id, filename, mime, file_bytes, owner_email=payload.get("email", ""))

    await _audit_log_lazy()(db, payload.get("email", ""), "tool_create_file", f"Created {req.file_type}: {filename}")
    await db.commit()
    return {"file_id": file_id, "filename": filename, "download_url": f"/api/ghost/tools/download/{file_id}", "size": len(file_bytes)}


# ─── ASYNC FILE CREATION (for mobile background support) ─────────────────

JOB_STORE: dict[str, dict] = {}  # job_id -> {status, file_id, error}

@router.post("/create-file-async")
async def create_file_async(req: CreateFileRequest, authorization: str = Header(...), db: AsyncSession = Depends(get_db)):
    """Start file creation in background. Returns job_id immediately."""
    payload = _verify_auth(authorization)
    await _audit_log_lazy()(db, payload.get("email", ""), "tool_create_file_async", f"Async {req.file_type} file")
    await db.commit()
    import asyncio
    job_id = str(uuid.uuid4())[:8]
    JOB_STORE[job_id] = {"status": "processing", "file_id": None, "error": None, "owner_email": payload.get("email", "")}

    async def _background_create():
        import time as _time
        t0 = _time.time()
        logger.info("[create-file-async] job=%s started type=%s desc='%s'", job_id, req.file_type, req.description[:100])
        try:
            # ── Check for structured document templates ──
            if req.file_type == "pdf":
                from lib.pdf_document_templates import detect_document_type, get_structured_prompt, render_structured_pdf
                doc_type = detect_document_type(req.description)
                if doc_type:
                    logger.info("[create-file-async] job=%s detected structured type=%s", job_id, doc_type)
                    structured_prompt = get_structured_prompt(doc_type)
                    raw_json = await _ask_claude(req.description, structured_prompt)
                    clean = raw_json.strip()
                    if clean.startswith('```'):
                        clean = clean.split('\n', 1)[1] if '\n' in clean else clean[3:]
                        if clean.endswith('```'):
                            clean = clean[:-3]
                        clean = clean.strip()
                    try:
                        data = json.loads(clean)
                        file_bytes = render_structured_pdf(doc_type, data)
                        file_id = str(uuid.uuid4())[:8]
                        filename = req.filename or f"{doc_type}_{file_id}"
                        filename += ".pdf"
                        mime = "application/pdf"
                        await _store_file(file_id, filename, mime, file_bytes)
                        elapsed = _time.time() - t0
                        logger.info("[create-file-async] job=%s DONE structured %s file=%s elapsed=%.1fs", job_id, doc_type, filename, elapsed)
                        JOB_STORE[job_id] = {"status": "done", "file_id": file_id, "filename": filename, "download_url": f"/api/ghost/tools/download/{file_id}", "size": len(file_bytes), "error": None}
                        return
                    except (json.JSONDecodeError, Exception) as e:
                        logger.warning("[create-file-async] job=%s structured %s failed (%s), falling back", job_id, doc_type, e)

            # Reuse the same logic as create_file
            format_instructions = {
                "csv": "Return raw CSV data with headers.",
                "txt": "Return well-written plain text with clear structure.",
                "xlsx": """Return CSV data (comma-separated values). First line is the title/report name. Second line is column headers. Remaining lines are data rows.
For calculated rows (totals, averages), put the actual calculated number, NOT a formula.
Use numbers for numeric values (no dollar signs, no quotes around numbers).
Example:
Monthly Budget Report
Category,Budget,Actual,Difference
Salary,5000,5000,0
Rent,1800,1800,0
TOTAL,6800,6800,0

IMPORTANT: Return ONLY the CSV data. No markdown fences. No explanations.""",
                "pdf": "Return well-structured text using markdown-style headings (# ## ###). Use - for bullet points. Use **bold** for emphasis. For financial documents, include tables using | pipes.",
                "docx": "Return well-structured text using markdown-style headings.",
            }

            # Detect accounting templates and enhance the description
            desc_lower = req.description.lower()
            accounting_hint = ""
            if any(term in desc_lower for term in ["p&l", "profit and loss", "profit & loss", "income statement"]):
                accounting_hint = "\nThis is a Profit & Loss statement. Include: Revenue, Cost of Goods Sold, Gross Profit, Operating Expenses (broken down), Operating Income, Net Income. Use formulas for all calculated rows."
            elif any(term in desc_lower for term in ["balance sheet"]):
                accounting_hint = "\nThis is a Balance Sheet. Include: Assets (Current + Non-current), Liabilities (Current + Long-term), Equity. Assets must equal Liabilities + Equity. Use formulas."
            elif any(term in desc_lower for term in ["expense report"]):
                accounting_hint = "\nThis is an Expense Report. Include: Date, Description, Category, Amount, Receipt columns. Add category subtotals and grand total with formulas."
            elif any(term in desc_lower for term in ["tax summary", "tax report", "deductible"]):
                accounting_hint = "\nThis is a Tax Summary. Group expenses by deductible category. Include totals per category and overall deductible total with formulas."
            elif any(term in desc_lower for term in ["invoice"]):
                accounting_hint = "\nThis is an Invoice. Include: Item, Description, Quantity, Unit Price, Amount. Add subtotal, tax, and total with formulas."

            system = f"""You are an expert professional writer and accountant. Create high-quality content.
DOCUMENT TYPE: {req.file_type.upper()}
FORMAT: {format_instructions.get(req.file_type, 'Return clean text.')}{accounting_hint}
Return ONLY the document content, no explanations."""

            logger.info("[create-file-async] job=%s calling Claude...", job_id)
            content = await _ask_claude(req.description, system)
            logger.info("[create-file-async] job=%s Claude done, content_len=%d elapsed=%.1fs", job_id, len(content), _time.time() - t0)
            file_id = str(uuid.uuid4())[:8]
            filename = req.filename or f"document_{file_id}"

            # Generate file bytes (same as sync endpoint)
            if req.file_type == "pdf":
                try:
                    from lib.pdf_templates import create_professional_pdf
                    file_bytes = create_professional_pdf(content, title=filename)
                    filename += ".pdf"
                    mime = "application/pdf"
                    logger.info("[create-file-async] job=%s PDF built, size=%d", job_id, len(file_bytes))
                except Exception as pdf_err:
                    logger.error("[create-file-async] job=%s PDF template failed: %s", job_id, pdf_err)
                    try:
                        from reportlab.lib.pagesizes import letter
                        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                        from reportlab.lib.styles import getSampleStyleSheet
                        buf = io.BytesIO()
                        doc = SimpleDocTemplate(buf, pagesize=letter)
                        styles = getSampleStyleSheet()
                        story = []
                        for line in content.split('\n'):
                            line = line.strip()
                            if not line:
                                story.append(Spacer(1, 12))
                            else:
                                clean = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
                                clean = re.sub(r'\*(.+?)\*', r'\1', clean)
                                clean = re.sub(r'^#+\s*', '', clean)
                                story.append(Paragraph(clean, styles['Normal']))
                        doc.build(story)
                        file_bytes = buf.getvalue()
                        filename += ".pdf"
                        mime = "application/pdf"
                    except Exception:
                        file_bytes = content.encode('utf-8')
                        filename += ".txt"
                        mime = "text/plain"
            elif req.file_type == "xlsx":
                try:
                    file_bytes, fname, mime = _csv_to_styled_xlsx(content, filename)
                    filename = fname
                except Exception as xlsx_err:
                    logger.error("XLSX builder failed: %s", xlsx_err)
                    file_bytes = content.encode('utf-8')
                    filename += ".csv"
                    mime = "text/csv"
            elif req.file_type == "docx":
                file_bytes, fname, mime = _content_to_file(content, "docx", filename)
                filename = fname
            elif req.file_type == "csv":
                file_bytes = content.encode('utf-8')
                filename += ".csv"
                mime = "text/csv"
            else:
                file_bytes = content.encode('utf-8')
                filename += ".txt"
                mime = "text/plain"

            await _store_file(file_id, filename, mime, file_bytes)
            elapsed = _time.time() - t0
            logger.info("[create-file-async] job=%s DONE file=%s size=%d elapsed=%.1fs", job_id, filename, len(file_bytes), elapsed)
            JOB_STORE[job_id] = {"status": "done", "file_id": file_id, "filename": filename, "download_url": f"/api/ghost/tools/download/{file_id}", "size": len(file_bytes), "error": None}
        except Exception as e:
            elapsed = _time.time() - t0
            logger.error("[create-file-async] job=%s FAILED after %.1fs: %s", job_id, elapsed, str(e))
            JOB_STORE[job_id] = {"status": "failed", "file_id": None, "error": str(e)}

    asyncio.create_task(_background_create())
    return {"job_id": job_id, "status": "processing"}


@router.get("/job-status/{job_id}")
async def job_status(job_id: str, authorization: str = Header(...)):
    """Check if an async file creation job is done — requires auth + ownership."""
    payload = _verify_auth(authorization)
    job = JOB_STORE.get(job_id)
    if not job:
        raise HTTPException(404, "Job not found")
    if job.get("owner_email") and job["owner_email"] != payload.get("email", ""):
        raise HTTPException(404, "Job not found")
    return {k: v for k, v in job.items() if k != "owner_email"}


@router.get("/download/{file_id}")
async def download_file(file_id: str, authorization: str = Header(...), db: AsyncSession = Depends(get_db)):
    """Download file — requires authentication and ownership check."""
    from fastapi.responses import Response
    payload = _verify_auth(authorization)

    f = await _get_file(file_id)
    if not f:
        raise HTTPException(404, "File not found or expired")
    # Ownership check — deny access if file belongs to a different user
    if f.get("owner_email") and f["owner_email"] != payload.get("email", ""):
        raise HTTPException(404, "File not found or expired")
    await _audit_log_lazy()(db, payload.get("email", ""), "tool_download_file", f"Downloaded {file_id}")
    await db.commit()
    return Response(content=base64.b64decode(f["data"]), media_type=f["mime"], headers={"Content-Disposition": f'attachment; filename="{f["filename"]}"'})


# ─── FILE MODIFICATION (edit, chart, convert, merge, filter) ──────────────

class ModifyFileRequest(BaseModel):
    operation: str  # edit, chart, convert, merge, filter, compare
    file_id: Optional[str] = None
    file_ids: Optional[list[str]] = None
    instructions: Optional[str] = ""
    target_format: Optional[str] = None  # for convert: pdf, xlsx, docx, csv, txt


@router.post("/modify-file-async")
async def modify_file_async(req: ModifyFileRequest, authorization: str = Header(...), db: AsyncSession = Depends(get_db)):
    """Modify an existing file: edit, chart, convert, merge, or filter."""
    payload = _verify_auth(authorization)
    await _audit_log_lazy()(db, payload.get("email", ""), "tool_modify_file", f"Operation: {req.operation}")
    await db.commit()
    import asyncio
    job_id = str(uuid.uuid4())[:8]
    JOB_STORE[job_id] = {"status": "processing", "file_id": None, "error": None, "owner_email": payload.get("email", "")}

    async def _background_modify():
        import time as _time
        t0 = _time.time()
        logger.info("[modify-file-async] job=%s started op=%s instructions='%s'", job_id, req.operation, (req.instructions or '')[:100])
        try:
            # Load original file(s) — check cache then R2
            owner = payload.get("email", "")
            if req.operation == "merge" and req.file_ids:
                sources = []
                for fid in req.file_ids:
                    f = await _get_file(fid)
                    if not f:
                        raise ValueError(f"File {fid} not found")
                    if f.get("owner_email") and f["owner_email"] != owner:
                        raise ValueError(f"File {fid} not found")
                    content = base64.b64decode(f["data"])
                    sources.append({"filename": f["filename"], "content": content, "mime": f["mime"]})
            elif req.file_id:
                f = await _get_file(req.file_id)
                if not f:
                    raise ValueError("File not found")
                if f.get("owner_email") and f["owner_email"] != owner:
                    raise ValueError("File not found")
                original_bytes = base64.b64decode(f["data"])
                original_filename = f["filename"]
                original_mime = f["mime"]
            else:
                raise ValueError("File not found")

            file_id = str(uuid.uuid4())[:8]
            # Strip previous operation suffixes to prevent name stacking
            base_name = original_filename.rsplit('.', 1)[0] if '.' in original_filename else original_filename
            for _suffix in ('_modified', '_converted', '_filtered', '_merged', '_chart', '_comparison'):
                while base_name.endswith(_suffix):
                    base_name = base_name[:-len(_suffix)]
            ext = original_filename.rsplit('.', 1)[-1] if '.' in original_filename else 'txt'

            if req.operation == "edit":
                is_spreadsheet = ext in ('xlsx', 'xls', 'csv')
                is_pdf = ext == 'pdf'

                # PDF page-level operations (delete, remove, reorder pages)
                if is_pdf and req.instructions and any(kw in req.instructions.lower() for kw in ['delete page', 'remove page', 'delete a page', 'remove a page', 'delete one page', 'remove one page', 'delete 1 page', 'remove 1 page']):
                    try:
                        import fitz  # PyMuPDF
                        pdf_doc = fitz.open(stream=original_bytes, filetype="pdf")
                        total_pages = len(pdf_doc)
                        if total_pages <= 1:
                            raise ValueError("Cannot delete — PDF only has 1 page")
                        # Find which page to delete (default: last page)
                        import re as _re
                        page_match = _re.search(r'page\s*(\d+)', req.instructions.lower())
                        if page_match:
                            page_num = int(page_match.group(1)) - 1  # 0-indexed
                        else:
                            page_num = total_pages - 1  # default: last page
                        page_num = max(0, min(page_num, total_pages - 1))
                        pdf_doc.delete_page(page_num)
                        buf = io.BytesIO()
                        pdf_doc.save(buf)
                        pdf_doc.close()
                        file_bytes = buf.getvalue()
                        filename = f"{base_name}_modified.pdf"
                        mime = "application/pdf"
                    except ImportError:
                        raise ValueError("PDF manipulation not available")
                elif is_spreadsheet:
                    # For spreadsheets: ask Claude to return JSON with formulas
                    system = """You are an Excel expert. The user has a spreadsheet and wants modifications.
Return a JSON object with this structure:
{"headers": ["col1", "col2", ...], "rows": [["val1", "val2", ...], ...], "formulas": {"C10": "=SUM(C2:C9)", "D10": "=AVERAGE(D2:D9)"}}

IMPORTANT:
- "rows" contains the DATA rows (not headers)
- "formulas" is a dict mapping cell references (like "C10") to Excel formula strings
- Use real Excel formulas (=SUM, =AVERAGE, =IF, =VLOOKUP, etc.) wherever appropriate
- For totals rows, running balances, calculated columns — ALWAYS use formulas, never static values
- Return ONLY valid JSON, no explanations."""
                    text_content = _extract_text(original_bytes, original_mime)
                    prompt = f"Current spreadsheet data:\n{text_content}\n\nModifications: {req.instructions}"
                    modified_content = await _ask_claude(prompt, system)
                    file_bytes, filename, mime = _content_to_xlsx_with_formulas(modified_content, f"{base_name}_modified")
                else:
                    text_content = _extract_text(original_bytes, original_mime)
                    system = """You are a document editor. The user has an existing document and wants modifications.
Return ONLY the complete modified document content. Keep the same format and structure unless told otherwise."""
                    prompt = f"Original document content:\n\n{text_content}\n\nModifications requested: {req.instructions}\n\nReturn the complete modified document."
                    modified_content = await _ask_claude(prompt, system)
                    file_bytes, filename, mime = _content_to_file(modified_content, ext, f"{base_name}_modified")

            elif req.operation == "chart":
                text_content = _extract_text(original_bytes, original_mime)
                system = """You are a data analyst. Given spreadsheet/CSV data, generate Python code that uses matplotlib to create the requested chart.
The code must:
1. Parse the data (provided as a string variable called DATA)
2. Create the chart using matplotlib
3. Save to a BytesIO buffer called 'buf'
4. Use plt.tight_layout() before saving
Return ONLY the Python code, no explanations."""
                prompt = f"Data:\n{text_content}\n\nCreate this chart: {req.instructions or 'Create an appropriate chart for this data'}"
                code = await _ask_claude(prompt, system)
                # Execute chart code in a sandboxed subprocess
                import subprocess, tempfile
                clean_code = code.replace("```python", "").replace("```", "").strip()

                # AST-based security check for chart code
                import ast as _chart_ast
                _CHART_ALLOWED = {"matplotlib", "math", "statistics", "json", "datetime", "re",
                                  "collections", "itertools", "functools", "decimal", "random",
                                  "string", "csv", "io", "sys", "numpy", "pandas"}
                _CHART_BLOCKED_BUILTINS = {"__import__", "exec", "eval", "compile", "open",
                                           "getattr", "setattr", "delattr", "globals", "locals",
                                           "vars", "dir", "type", "breakpoint", "input", "help"}
                try:
                    chart_tree = _chart_ast.parse(clean_code)
                    for cnode in _chart_ast.walk(chart_tree):
                        if isinstance(cnode, _chart_ast.Import):
                            for alias in cnode.names:
                                mod = alias.name.split(".")[0]
                                if mod not in _CHART_ALLOWED:
                                    raise HTTPException(400, f"Chart code: blocked import '{mod}'")
                        elif isinstance(cnode, _chart_ast.ImportFrom) and cnode.module:
                            mod = cnode.module.split(".")[0]
                            if mod not in _CHART_ALLOWED:
                                raise HTTPException(400, f"Chart code: blocked import '{mod}'")
                        elif isinstance(cnode, _chart_ast.Call):
                            fn = cnode.func
                            if isinstance(fn, _chart_ast.Name) and fn.id in _CHART_BLOCKED_BUILTINS:
                                raise HTTPException(400, f"Chart code: blocked '{fn.id}()'")
                            elif isinstance(fn, _chart_ast.Attribute) and fn.attr in _CHART_BLOCKED_BUILTINS:
                                raise HTTPException(400, f"Chart code: blocked '.{fn.attr}()'")
                        elif isinstance(cnode, _chart_ast.Attribute) and cnode.attr.startswith("__") and cnode.attr.endswith("__"):
                            raise HTTPException(400, f"Chart code: blocked dunder '{cnode.attr}'")
                except SyntaxError as se:
                    raise HTTPException(400, f"Chart code syntax error: {se}")

                # Write a self-contained script that outputs PNG to stdout
                with tempfile.NamedTemporaryFile(mode="w", suffix=".py", delete=False) as tmp:
                    tmp.write("import matplotlib\nmatplotlib.use('Agg')\nimport matplotlib.pyplot as plt\nimport io, sys\n")
                    tmp.write(f"DATA = {repr(text_content)}\n")
                    tmp.write("buf = io.BytesIO()\n")
                    tmp.write(clean_code + "\n")
                    tmp.write("buf.seek(0)\nsys.stdout.buffer.write(buf.getvalue())\n")
                    tmp_path = tmp.name

                try:
                    result = subprocess.run(
                        ["python3", tmp_path],
                        capture_output=True, timeout=15,
                        env={"PATH": "/usr/bin:/usr/local/bin", "HOME": "/tmp"},
                    )
                    if result.returncode != 0:
                        logger.error("Chart code failed: %s", result.stderr.decode()[:500])
                        raise HTTPException(400, "Chart generation failed. Try a different chart type.")
                    file_bytes = result.stdout
                    if not file_bytes:
                        raise HTTPException(400, "Chart generation produced no output.")
                finally:
                    import os as _os
                    _os.unlink(tmp_path)

                filename = f"{base_name}_chart.png"
                mime = "image/png"

            elif req.operation == "convert":
                target = req.target_format or "pdf"
                file_bytes, filename, mime = _convert_file(original_bytes, original_mime, ext, target, f"{base_name}_converted")

            elif req.operation == "merge":
                combined_text = ""
                for src in sources:
                    text = _extract_text(src["content"], src["mime"])
                    combined_text += f"\n\n--- {src['filename']} ---\n\n{text}"
                # Determine output format from first file
                ext = sources[0]["filename"].rsplit('.', 1)[-1] if '.' in sources[0]["filename"] else 'txt'
                file_bytes, filename, mime = _content_to_file(combined_text, ext, f"{base_name}_merged")

            elif req.operation == "filter":
                text_content = _extract_text(original_bytes, original_mime)
                system = """You are a data processor. Filter the data according to the user's criteria.
Return ONLY the filtered data in the same format (CSV for CSV, JSON array for XLSX)."""
                prompt = f"Data:\n{text_content}\n\nFilter criteria: {req.instructions}"
                filtered = await _ask_claude(prompt, system)
                ext = original_filename.rsplit('.', 1)[-1] if '.' in original_filename else 'csv'
                file_bytes, filename, mime = _content_to_file(filtered, ext, f"{base_name}_filtered")

            elif req.operation == "compare":
                if not req.file_ids or len(req.file_ids) < 2:
                    raise ValueError("Compare requires at least 2 file IDs")
                texts = []
                for fid in req.file_ids[:2]:
                    f = await _get_file(fid)
                    if not f:
                        raise ValueError(f"File {fid} not found")
                    if f.get("owner_email") and f["owner_email"] != owner:
                        raise ValueError(f"File {fid} not found")
                    content = base64.b64decode(f["data"])
                    texts.append({"name": f["filename"], "content": _extract_text(content, f["mime"])})
                system = """You are a spreadsheet analyst. Compare two spreadsheets and produce a detailed comparison report.
For each difference found, show: row/column location, old value, new value.
Group changes by type: Added rows, Removed rows, Modified values, New columns, Removed columns.
Use markdown formatting with tables where appropriate. Be thorough but concise."""
                prompt = f"File 1 ({texts[0]['name']}):\n{texts[0]['content']}\n\nFile 2 ({texts[1]['name']}):\n{texts[1]['content']}\n\n{req.instructions or 'Compare these files and highlight all differences.'}"
                comparison = await _ask_claude(prompt, system)
                try:
                    from lib.pdf_templates import create_professional_pdf
                    file_bytes = create_professional_pdf(comparison, title=f"Comparison Report")
                except Exception:
                    file_bytes = comparison.encode('utf-8')
                filename = f"{base_name}_comparison.pdf"
                mime = "application/pdf"

            elif req.operation == "reconcile":
                if not req.file_ids or len(req.file_ids) < 2:
                    raise ValueError("Reconcile requires 2 file IDs (bank statement + your records)")
                texts = []
                for fid in req.file_ids[:2]:
                    f = await _get_file(fid)
                    if not f:
                        raise ValueError(f"File {fid} not found")
                    if f.get("owner_email") and f["owner_email"] != owner:
                        raise ValueError(f"File {fid} not found")
                    content = base64.b64decode(f["data"])
                    texts.append({"name": f["filename"], "content": _extract_text(content, f["mime"])})

                system = """You are an expert accountant performing a bank reconciliation.
You have two data sources:
1. BANK STATEMENT — official transactions from the bank
2. BOOK RECORDS — the user's own accounting records

Your job:
1. Match transactions between the two sources by amount AND approximate date (±3 days). Descriptions may differ (e.g. "AMZN*123" vs "Amazon supplies" are the same).
2. Classify every transaction into: MATCHED, BANK ONLY (in bank but not in books), BOOKS ONLY (in books but not in bank).

Return a JSON object with this EXACT structure:
{
  "matched": [{"bank_date": "...", "bank_desc": "...", "book_date": "...", "book_desc": "...", "amount": 0.00}],
  "bank_only": [{"date": "...", "description": "...", "amount": 0.00}],
  "books_only": [{"date": "...", "description": "...", "amount": 0.00}],
  "bank_total": 0.00,
  "books_total": 0.00,
  "difference": 0.00,
  "matched_count": 0,
  "unmatched_count": 0
}

Be thorough. Match fuzzy descriptions. Return ONLY valid JSON."""

                prompt = f"BANK STATEMENT ({texts[0]['name']}):\n{texts[0]['content']}\n\nBOOK RECORDS ({texts[1]['name']}):\n{texts[1]['content']}\n\n{req.instructions or 'Reconcile these two sources.'}"
                result_json = await _ask_claude(prompt, system)

                # Build a styled Excel reconciliation report
                try:
                    import json as _json
                    from openpyxl import Workbook
                    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

                    data = _json.loads(result_json)
                    wb = Workbook()

                    thin_border = Border(
                        left=Side(style='thin', color='D9D9D9'), right=Side(style='thin', color='D9D9D9'),
                        top=Side(style='thin', color='D9D9D9'), bottom=Side(style='thin', color='D9D9D9'),
                    )
                    header_font = Font(bold=True, color="FFFFFF", size=11)

                    # Summary sheet
                    ws_summary = wb.active
                    ws_summary.title = "Summary"
                    summary_fill = PatternFill(start_color="2B579A", end_color="2B579A", fill_type="solid")
                    for r, (label, value) in enumerate([
                        ("Bank Statement Total", data.get("bank_total", 0)),
                        ("Book Records Total", data.get("books_total", 0)),
                        ("Difference", data.get("difference", 0)),
                        ("", ""),
                        ("Matched Transactions", data.get("matched_count", len(data.get("matched", [])))),
                        ("Bank Only (not in books)", len(data.get("bank_only", []))),
                        ("Books Only (not in bank)", len(data.get("books_only", []))),
                    ], 1):
                        ws_summary.cell(row=r, column=1, value=label).font = Font(bold=True, size=12)
                        cell = ws_summary.cell(row=r, column=2, value=value)
                        if label == "Difference" and value != 0:
                            cell.font = Font(bold=True, color="FF0000", size=12)
                        else:
                            cell.font = Font(size=12)
                    ws_summary.column_dimensions['A'].width = 30
                    ws_summary.column_dimensions['B'].width = 20

                    # Matched sheet (green)
                    ws_matched = wb.create_sheet("Matched")
                    green_fill = PatternFill(start_color="22C55E", end_color="22C55E", fill_type="solid")
                    headers = ["Bank Date", "Bank Description", "Book Date", "Book Description", "Amount"]
                    for c, h in enumerate(headers, 1):
                        cell = ws_matched.cell(row=1, column=c, value=h)
                        cell.font = header_font
                        cell.fill = green_fill
                        cell.border = thin_border
                    for r, item in enumerate(data.get("matched", []), 2):
                        ws_matched.cell(row=r, column=1, value=item.get("bank_date", "")).border = thin_border
                        ws_matched.cell(row=r, column=2, value=item.get("bank_desc", "")).border = thin_border
                        ws_matched.cell(row=r, column=3, value=item.get("book_date", "")).border = thin_border
                        ws_matched.cell(row=r, column=4, value=item.get("book_desc", "")).border = thin_border
                        ws_matched.cell(row=r, column=5, value=item.get("amount", 0)).border = thin_border
                    for col in ws_matched.columns:
                        ws_matched.column_dimensions[col[0].column_letter].width = 25

                    # Bank Only sheet (red)
                    ws_bank = wb.create_sheet("Bank Only")
                    red_fill = PatternFill(start_color="EF4444", end_color="EF4444", fill_type="solid")
                    for c, h in enumerate(["Date", "Description", "Amount"], 1):
                        cell = ws_bank.cell(row=1, column=c, value=h)
                        cell.font = header_font
                        cell.fill = red_fill
                        cell.border = thin_border
                    for r, item in enumerate(data.get("bank_only", []), 2):
                        ws_bank.cell(row=r, column=1, value=item.get("date", "")).border = thin_border
                        ws_bank.cell(row=r, column=2, value=item.get("description", "")).border = thin_border
                        ws_bank.cell(row=r, column=3, value=item.get("amount", 0)).border = thin_border
                    for col in ws_bank.columns:
                        ws_bank.column_dimensions[col[0].column_letter].width = 25

                    # Books Only sheet (orange)
                    ws_books = wb.create_sheet("Books Only")
                    orange_fill = PatternFill(start_color="F59E0B", end_color="F59E0B", fill_type="solid")
                    for c, h in enumerate(["Date", "Description", "Amount"], 1):
                        cell = ws_books.cell(row=1, column=c, value=h)
                        cell.font = header_font
                        cell.fill = orange_fill
                        cell.border = thin_border
                    for r, item in enumerate(data.get("books_only", []), 2):
                        ws_books.cell(row=r, column=1, value=item.get("date", "")).border = thin_border
                        ws_books.cell(row=r, column=2, value=item.get("description", "")).border = thin_border
                        ws_books.cell(row=r, column=3, value=item.get("amount", 0)).border = thin_border
                    for col in ws_books.columns:
                        ws_books.column_dimensions[col[0].column_letter].width = 25

                    buf = io.BytesIO()
                    wb.save(buf)
                    file_bytes = buf.getvalue()
                    filename = f"reconciliation_{file_id}.xlsx"
                    mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                except Exception:
                    # Fallback to PDF
                    try:
                        from lib.pdf_templates import create_professional_pdf
                        file_bytes = create_professional_pdf(result_json, title="Reconciliation Report")
                    except Exception:
                        file_bytes = result_json.encode('utf-8')
                    filename = f"reconciliation_{file_id}.pdf"
                    mime = "application/pdf"

            else:
                raise ValueError(f"Unknown operation: {req.operation}")

            await _store_file(file_id, filename, mime, file_bytes)
            elapsed = _time.time() - t0
            logger.info("[modify-file-async] job=%s DONE file=%s size=%d elapsed=%.1fs", job_id, filename, len(file_bytes), elapsed)
            JOB_STORE[job_id] = {"status": "done", "file_id": file_id, "filename": filename, "download_url": f"/api/ghost/tools/download/{file_id}", "size": len(file_bytes), "error": None}
        except Exception as e:
            elapsed = _time.time() - t0
            logger.error("[modify-file-async] job=%s FAILED after %.1fs: %s", job_id, elapsed, str(e))
            JOB_STORE[job_id] = {"status": "failed", "file_id": None, "error": str(e)}

    asyncio.create_task(_background_modify())
    return {"job_id": job_id, "status": "processing"}


def _csv_to_styled_xlsx(content: str, base_name: str) -> tuple:
    """Convert CSV content from Claude into a professional styled Excel file."""
    import csv as _csv
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    from datetime import datetime

    # Clean content — remove markdown fences
    import re
    cleaned = re.sub(r'^```\w*\s*\n?', '', content.strip())
    cleaned = re.sub(r'\n?```\s*$', '', cleaned).strip()

    # Parse lines
    lines = [l for l in cleaned.split('\n') if l.strip()]
    if not lines:
        raise ValueError("No content")

    # First line is the title (if it doesn't contain commas or has fewer than the header line)
    title = base_name.replace("_", " ").title()
    data_lines = lines

    # Detect title: first line with fewer commas than second line
    if len(lines) >= 2:
        first_commas = lines[0].count(',')
        second_commas = lines[1].count(',')
        if first_commas < second_commas or first_commas == 0:
            title = lines[0].strip()
            data_lines = lines[1:]

    # Parse CSV
    reader = _csv.reader(data_lines)
    all_rows = list(reader)
    if not all_rows:
        raise ValueError("No data rows")

    headers = all_rows[0]
    data_rows = all_rows[1:]

    # Create workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Report"
    num_cols = len(headers)
    last_col = get_column_letter(num_cols) if num_cols > 0 else 'A'

    # Styles
    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill(start_color="2B579A", end_color="2B579A", fill_type="solid")
    title_font = Font(bold=True, size=16, color="2B579A")
    date_font = Font(size=10, color="888888", italic=True)
    border = Border(
        left=Side(style='thin', color='D0D0D0'), right=Side(style='thin', color='D0D0D0'),
        top=Side(style='thin', color='D0D0D0'), bottom=Side(style='thin', color='D0D0D0'),
    )
    alt_fill = PatternFill(start_color="F7F9FC", end_color="F7F9FC", fill_type="solid")
    total_fill = PatternFill(start_color="E8EDF5", end_color="E8EDF5", fill_type="solid")

    # Row 1: Title
    ws.merge_cells(f'A1:{last_col}1')
    ws['A1'].value = title
    ws['A1'].font = title_font
    ws['A1'].alignment = Alignment(horizontal='left', vertical='center')
    ws.row_dimensions[1].height = 30

    # Row 2: Date
    ws.merge_cells(f'A2:{last_col}2')
    ws['A2'].value = f"Generated {datetime.utcnow().strftime('%B %d, %Y')}"
    ws['A2'].font = date_font
    ws.row_dimensions[2].height = 18

    # Row 3: Spacer
    ws.row_dimensions[3].height = 6

    # Row 4: Headers
    for col_idx, h in enumerate(headers, 1):
        cell = ws.cell(row=4, column=col_idx, value=h.strip())
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = border
    ws.row_dimensions[4].height = 28

    # Row 5+: Data
    for row_idx, row in enumerate(data_rows):
        excel_row = row_idx + 5
        is_total = False

        for col_idx, val in enumerate(row):
            if col_idx >= num_cols:
                break
            cell = ws.cell(row=excel_row, column=col_idx + 1)
            val = val.strip()

            # Try to convert to number
            try:
                clean_val = val.replace(',', '').replace('$', '').replace('%', '').strip()
                if '.' in clean_val:
                    cell.value = float(clean_val)
                    cell.number_format = '#,##0.00'
                elif clean_val.lstrip('-').isdigit():
                    cell.value = int(clean_val)
                    cell.number_format = '#,##0'
                else:
                    cell.value = val
            except (ValueError, AttributeError):
                cell.value = val

            cell.border = border

            # Check if total row
            if col_idx == 0 and isinstance(cell.value, str):
                lower = cell.value.lower()
                if any(kw in lower for kw in ['total', 'gross', 'net', 'subtotal', 'balance']):
                    is_total = True

            # Number alignment
            if isinstance(cell.value, (int, float)):
                cell.alignment = Alignment(horizontal='right', vertical='center')
                if cell.value < 0:
                    cell.font = Font(color="CC0000")

        # Row styling
        for col_idx in range(num_cols):
            cell = ws.cell(row=excel_row, column=col_idx + 1)
            if is_total:
                cell.fill = total_fill
                cell.font = Font(bold=True, size=11) if not (isinstance(cell.value, (int, float)) and cell.value < 0) else Font(bold=True, color="CC0000", size=11)
            elif row_idx % 2 == 1:
                cell.fill = alt_fill

            # Section headers (ALL CAPS first col)
            if col_idx == 0 and isinstance(cell.value, str) and cell.value.isupper() and not is_total:
                cell.font = Font(bold=True, size=11, color="2B579A")

    # Freeze headers
    ws.freeze_panes = 'A5'

    # Auto-width
    for col_idx in range(1, num_cols + 1):
        col_letter = get_column_letter(col_idx)
        max_len = 0
        for row in ws.iter_rows(min_row=4, max_row=ws.max_row, min_col=col_idx, max_col=col_idx):
            for cell in row:
                if cell.value is not None:
                    if isinstance(cell.value, (int, float)):
                        val_len = len(f"{cell.value:,}")
                    else:
                        val_len = len(str(cell.value))
                    max_len = max(max_len, val_len)
        if col_idx == 1:
            ws.column_dimensions[col_letter].width = max(min(max_len + 4, 35), 20)
        else:
            ws.column_dimensions[col_letter].width = max(min(max_len + 4, 22), 14)

    # Print setup
    ws.page_setup.orientation = 'landscape'
    ws.page_setup.fitToWidth = 1

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue(), f"{base_name}.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"


def _extract_json_from_content(content: str):
    """Extract JSON from Claude's response, handling all possible formats."""
    import json as _json
    import re

    cleaned = content.strip()

    # Remove markdown fences aggressively
    cleaned = re.sub(r'^```\w*\s*\n?', '', cleaned)
    cleaned = re.sub(r'\n?```\s*$', '', cleaned)
    cleaned = cleaned.strip()

    # Find the outermost JSON object or array
    # Track brace depth to find matching close
    start = -1
    for i, ch in enumerate(cleaned):
        if ch in ('{', '['):
            start = i
            break

    if start < 0:
        raise ValueError("No JSON found in content")

    open_ch = cleaned[start]
    close_ch = '}' if open_ch == '{' else ']'
    depth = 0
    end = -1
    in_string = False
    escape = False

    for i in range(start, len(cleaned)):
        ch = cleaned[i]
        if escape:
            escape = False
            continue
        if ch == '\\':
            escape = True
            continue
        if ch == '"':
            in_string = not in_string
            continue
        if in_string:
            continue
        if ch == open_ch:
            depth += 1
        elif ch == close_ch:
            depth -= 1
            if depth == 0:
                end = i
                break

    if end < 0:
        raise ValueError("Unbalanced JSON")

    json_str = cleaned[start:end + 1]
    return _json.loads(json_str)


def _content_to_xlsx_with_formulas(content: str, base_name: str) -> tuple:
    """Convert Claude's JSON response to Excel with real formulas."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

        data = _extract_json_from_content(content)
        logger.debug("XLSX PARSED JSON type=%s", type(data).__name__)

        # Handle both {"headers":[], "rows":[]} and [{"col":"val"}] formats
        if isinstance(data, list):
            headers = list(data[0].keys()) if data else []
            rows = [[row.get(h, "") for h in headers] for row in data]
            formulas = {}
            title = base_name.replace("_", " ").title()
        else:
            headers = data.get("headers", [])
            rows = data.get("rows", [])
            formulas = data.get("formulas", {})
            title = data.get("title", base_name.replace("_", " ").title())

        # Ensure headers and rows are clean lists
        if not isinstance(headers, list):
            headers = []
        if not isinstance(rows, list):
            rows = []
        # Keep all rows that are lists (don't filter — empty cells are normal for formula rows)
        rows = [r for r in rows if isinstance(r, list)]

        logger.debug("XLSX DATA: %d headers, %d rows, %d formulas", len(headers), len(rows), len(formulas))

        wb = Workbook()
        ws = wb.active
        ws.title = "Report"
        from openpyxl.utils import get_column_letter
        from datetime import datetime

        num_cols = max(len(headers), max((len(r) for r in rows), default=1)) if rows else len(headers)
        last_col = get_column_letter(num_cols) if num_cols > 0 else 'A'

        # --- Title section ---
        ws.merge_cells(f'A1:{last_col}1')
        title_cell = ws['A1']
        title_cell.value = title
        title_cell.font = Font(bold=True, size=16, color="2B579A")
        title_cell.alignment = Alignment(horizontal='left', vertical='center')
        ws.row_dimensions[1].height = 30

        ws.merge_cells(f'A2:{last_col}2')
        date_cell = ws['A2']
        date_cell.value = f"Generated {datetime.utcnow().strftime('%B %d, %Y')}"
        date_cell.font = Font(size=10, color="888888", italic=True)
        date_cell.alignment = Alignment(horizontal='left')
        ws.row_dimensions[2].height = 18

        # Empty spacer row
        ws.row_dimensions[3].height = 8

        # --- Styles ---
        header_font = Font(bold=True, color="FFFFFF", size=11)
        header_fill = PatternFill(start_color="2B579A", end_color="2B579A", fill_type="solid")
        thin_border = Border(
            left=Side(style='thin', color='D0D0D0'),
            right=Side(style='thin', color='D0D0D0'),
            top=Side(style='thin', color='D0D0D0'),
            bottom=Side(style='thin', color='D0D0D0'),
        )
        bottom_border = Border(bottom=Side(style='medium', color='2B579A'))
        alt_fill = PatternFill(start_color="F7F9FC", end_color="F7F9FC", fill_type="solid")
        total_fill = PatternFill(start_color="E8EDF5", end_color="E8EDF5", fill_type="solid")

        header_row = 4

        # --- Write headers ---
        if headers:
            for col_idx, h in enumerate(headers, 1):
                cell = ws.cell(row=header_row, column=col_idx, value=h)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = Alignment(horizontal='center', vertical='center')
                cell.border = thin_border
            ws.row_dimensions[header_row].height = 28

        # --- Write data rows ---
        data_start = header_row + 1
        for i, row in enumerate(rows):
            converted = []
            for val in row:
                if isinstance(val, str):
                    try:
                        cleaned_val = val.replace(',', '').replace('$', '').strip()
                        if '.' in cleaned_val: converted.append(float(cleaned_val))
                        else: converted.append(int(cleaned_val))
                    except (ValueError, TypeError):
                        converted.append(val)
                else:
                    converted.append(val)
            for col_idx, val in enumerate(converted, 1):
                ws.cell(row=data_start + i, column=col_idx, value=val)

        # --- Style data cells ---
        for row_idx, row in enumerate(ws.iter_rows(min_row=data_start, max_row=ws.max_row, max_col=num_cols)):
            is_total_row = False
            first_val = row[0].value if row[0].value else ""
            if isinstance(first_val, str):
                lower = first_val.lower()
                is_total_row = any(kw in lower for kw in ['total', 'gross', 'net', 'subtotal'])

            for cell in row:
                cell.border = thin_border
                cell.alignment = Alignment(vertical='center')

                # Number formatting
                if isinstance(cell.value, (int, float)):
                    cell.number_format = '#,##0' if isinstance(cell.value, int) else '#,##0.00'
                    cell.alignment = Alignment(horizontal='right', vertical='center')
                elif isinstance(cell.value, str) and len(cell.value) > 30:
                    cell.alignment = Alignment(wrap_text=True, vertical='center')
                    # Red for negative
                    if cell.value < 0:
                        cell.font = Font(color="CC0000")

                # Total/summary rows
                if is_total_row:
                    cell.fill = total_fill
                    cell.font = Font(bold=True, size=11)
                    if isinstance(cell.value, (int, float)) and cell.value < 0:
                        cell.font = Font(bold=True, color="CC0000", size=11)
                elif row_idx % 2 == 1:
                    cell.fill = alt_fill

                # Section headers (ALL CAPS in first column)
                if cell.column == 1 and isinstance(cell.value, str) and cell.value.isupper() and not is_total_row:
                    cell.font = Font(bold=True, size=11, color="2B579A")

        # --- Insert formulas ---
        # Adjust formula cell references to account for title rows (shift down by 3)
        for cell_ref, formula in formulas.items():
            # Adjust the row numbers in formulas
            import re
            col_letter = re.match(r'([A-Z]+)', cell_ref).group(1) if re.match(r'([A-Z]+)', cell_ref) else cell_ref[0]
            row_num = int(re.search(r'(\d+)', cell_ref).group(1)) if re.search(r'(\d+)', cell_ref) else 1
            new_ref = f"{col_letter}{row_num + 3}"
            # Also adjust references inside the formula
            def shift_ref(m):
                return f"{m.group(1)}{int(m.group(2)) + 3}"
            adjusted_formula = re.sub(r'([A-Z]+)(\d+)', shift_ref, formula)
            ws[new_ref] = adjusted_formula
            ws[new_ref].font = Font(bold=True, color="2B579A")
            ws[new_ref].border = thin_border
            ws[new_ref].number_format = '#,##0'
            ws[new_ref].fill = total_fill

        # --- Freeze panes (freeze headers) ---
        ws.freeze_panes = f'A{header_row + 1}'

        # --- Auto-width columns (generous padding to prevent overlap) ---
        for col in ws.columns:
            max_len = 0
            col_letter = col[0].column_letter
            for cell in col:
                try:
                    if cell.value is None:
                        continue
                    if isinstance(cell.value, (int, float)):
                        # Format numbers as they'll appear with commas
                        val = f"{cell.value:,.2f}" if isinstance(cell.value, float) else f"{cell.value:,}"
                    elif isinstance(cell.value, str) and cell.value.startswith('='):
                        # Formulas — estimate result width
                        val = "000,000,000"
                    else:
                        val = str(cell.value)
                    # Account for bold text being wider
                    width = len(val)
                    if cell.font and cell.font.bold:
                        width = int(width * 1.15)
                    max_len = max(max_len, width)
                except Exception:
                    pass
            # First column (labels) gets extra space, number columns get padding for commas
            if col_letter == 'A':
                ws.column_dimensions[col_letter].width = max(min(max_len + 6, 40), 18)
            else:
                ws.column_dimensions[col_letter].width = max(min(max_len + 4, 20), 14)

        # --- Print setup ---
        ws.sheet_properties.pageSetUpPr = None
        ws.page_setup.orientation = 'landscape'
        ws.page_setup.fitToWidth = 1
        ws.page_setup.fitToHeight = 0

        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue(), f"{base_name}.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    except Exception as e:
        logger.error("XLSX formula builder failed: %s", e)
        # Fallback: try to make a decent Excel from whatever content we got
        try:
            from openpyxl import Workbook as _FbWb
            from openpyxl.styles import Font as _FbFont, PatternFill as _FbFill, Border as _FbBorder, Side as _FbSide, Alignment as _FbAlign
            _wb = _FbWb()
            _ws = _wb.active
            _hfont = _FbFont(bold=True, color="FFFFFF", size=11)
            _hfill = _FbFill(start_color="2B579A", end_color="2B579A", fill_type="solid")
            _border = _FbBorder(left=_FbSide(style='thin', color='D0D0D0'), right=_FbSide(style='thin', color='D0D0D0'), top=_FbSide(style='thin', color='D0D0D0'), bottom=_FbSide(style='thin', color='D0D0D0'))
            # Try CSV-style parsing
            lines = [l.strip() for l in content.split('\n') if l.strip() and not l.strip().startswith('{') and not l.strip().startswith('[') and not l.strip().startswith('"') and not l.strip().startswith('}') and not l.strip().startswith(']')]
            if not lines:
                lines = [l.strip() for l in content.split('\n') if l.strip()]
            for row_idx, line in enumerate(lines):
                cells = [c.strip().strip('"').strip("'") for c in line.split(',')]
                for col_idx, val in enumerate(cells):
                    cell = _ws.cell(row=row_idx + 1, column=col_idx + 1)
                    try:
                        cell.value = int(val) if val.replace('-','').isdigit() else float(val) if val.replace('-','').replace('.','').isdigit() else val
                    except (ValueError, AttributeError):
                        cell.value = val
                    cell.border = _border
                    if row_idx == 0:
                        cell.font = _hfont
                        cell.fill = _hfill
                        cell.alignment = _FbAlign(horizontal='center')
                    elif isinstance(cell.value, (int, float)):
                        cell.number_format = '#,##0'
            for col in _ws.columns:
                _ws.column_dimensions[col[0].column_letter].width = 15
            _buf = io.BytesIO()
            _wb.save(_buf)
            return _buf.getvalue(), f"{base_name}.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        except Exception:
            return content.encode('utf-8'), f"{base_name}.txt", "text/plain"


def _convert_file(original_bytes: bytes, original_mime: str, source_ext: str, target_ext: str, base_name: str) -> tuple:
    """Convert between file formats using LibreOffice headless for quality,
    with Python-native fallbacks for simple cases."""
    import subprocess
    MIME_MAP = {
        "pdf": "application/pdf",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "csv": "text/csv",
        "txt": "text/plain",
        "pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
    }
    # LibreOffice filter names for target formats
    LO_FILTERS = {
        "pdf": "writer_pdf_Export",
        "docx": "MS Word 2007 XML",
        "xlsx": "Calc MS Excel 2007 XML",
        "csv": "Text - txt - csv (StarCalc)",
        "txt": "Text",
    }
    # Source extensions that LibreOffice handles well
    LO_SUPPORTED_SOURCE = {"pdf", "docx", "doc", "xlsx", "xls", "csv", "txt", "pptx", "ppt", "odt", "ods", "odp", "rtf"}

    # ── XLSX → CSV: use Python (faster, more reliable for this) ──
    if source_ext in ("xlsx", "xls") and target_ext == "csv":
        try:
            from openpyxl import load_workbook
            import csv as _csv
            wb = load_workbook(io.BytesIO(original_bytes))
            ws = wb.active
            buf = io.StringIO()
            writer = _csv.writer(buf)
            for row in ws.iter_rows(values_only=True):
                writer.writerow([c if c is not None else "" for c in row])
            return buf.getvalue().encode('utf-8'), f"{base_name}.csv", MIME_MAP["csv"]
        except Exception as e:
            logger.error("XLSX→CSV Python failed: %s", e)

    # ── CSV → XLSX: use Python (styled output) ──
    if source_ext == "csv" and target_ext == "xlsx":
        try:
            text = original_bytes.decode('utf-8', errors='ignore')
            file_bytes, fname, mime = _csv_to_styled_xlsx(text, base_name)
            return file_bytes, fname, mime
        except Exception as e:
            logger.error("CSV→XLSX Python failed: %s", e)

    # ── Any → TXT: just extract text ──
    if target_ext == "txt":
        text = _extract_text(original_bytes, original_mime)
        return text.encode('utf-8'), f"{base_name}.txt", MIME_MAP["txt"]

    # ── LibreOffice headless conversion (handles everything else) ──
    if source_ext in LO_SUPPORTED_SOURCE and target_ext in LO_FILTERS:
        try:
            with tempfile.TemporaryDirectory() as tmpdir:
                src_path = os.path.join(tmpdir, f"input.{source_ext}")
                with open(src_path, 'wb') as f:
                    f.write(original_bytes)

                # Pick the right filter for the target
                lo_filter = LO_FILTERS[target_ext]
                # For spreadsheet sources converting to PDF, use Calc filter
                if source_ext in ("xlsx", "xls", "csv") and target_ext == "pdf":
                    lo_filter = "calc_pdf_Export"
                # For presentations converting to PDF
                if source_ext in ("pptx", "ppt") and target_ext == "pdf":
                    lo_filter = "impress_pdf_Export"

                cmd = [
                    "libreoffice", "--headless", "--norestore", "--convert-to",
                    f"{target_ext}:{lo_filter}" if lo_filter else target_ext,
                    "--outdir", tmpdir, src_path
                ]
                logger.info("[convert] running: %s", " ".join(cmd))
                result = subprocess.run(cmd, capture_output=True, timeout=60, env={
                    "HOME": tmpdir, "PATH": "/usr/bin:/usr/local/bin:/bin",
                })
                if result.returncode != 0:
                    logger.error("[convert] LibreOffice failed: %s", result.stderr.decode()[:500])
                    raise RuntimeError("LibreOffice conversion failed")

                # Find the output file
                out_path = os.path.join(tmpdir, f"input.{target_ext}")
                if not os.path.exists(out_path):
                    # Sometimes LO uses a different name
                    for f in os.listdir(tmpdir):
                        if f.endswith(f".{target_ext}") and f != f"input.{source_ext}":
                            out_path = os.path.join(tmpdir, f)
                            break

                if not os.path.exists(out_path):
                    raise RuntimeError(f"Output file not found after conversion")

                with open(out_path, 'rb') as f:
                    file_bytes = f.read()

                filename = f"{base_name}.{target_ext}"
                mime = MIME_MAP.get(target_ext, "application/octet-stream")
                logger.info("[convert] LibreOffice success: %s → %s (%d bytes)", source_ext, target_ext, len(file_bytes))
                return file_bytes, filename, mime
        except subprocess.TimeoutExpired:
            logger.error("[convert] LibreOffice timed out after 60s")
        except Exception as e:
            logger.error("[convert] LibreOffice failed: %s", e)

    # ── Fallback: extract text and rebuild with Python ──
    logger.warning("[convert] falling back to text extraction: %s → %s", source_ext, target_ext)
    text = _extract_text(original_bytes, original_mime)
    return _content_to_file(text, target_ext, base_name)


def _extract_text(file_bytes: bytes, mime: str) -> str:
    """Extract readable text from file bytes based on mime type."""
    if mime == "application/pdf":
        try:
            import fitz
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            return "\n".join(page.get_text() for page in doc)
        except Exception:
            return file_bytes.decode('utf-8', errors='ignore')
    elif mime in ("application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "application/vnd.ms-excel"):
        try:
            from openpyxl import load_workbook
            wb = load_workbook(io.BytesIO(file_bytes))
            ws = wb.active
            rows = []
            for row in ws.iter_rows(values_only=True):
                rows.append(",".join(str(c) if c is not None else "" for c in row))
            return "\n".join(rows)
        except Exception:
            return file_bytes.decode('utf-8', errors='ignore')
    elif mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception:
            return file_bytes.decode('utf-8', errors='ignore')
    else:
        return file_bytes.decode('utf-8', errors='ignore')


def _content_to_file(content: str, ext: str, base_name: str) -> tuple:
    """Convert text content to file bytes in the specified format."""
    if ext == "pdf":
        try:
            from lib.pdf_templates import create_professional_pdf
            file_bytes = create_professional_pdf(content, title=base_name)
        except Exception:
            try:
                from reportlab.lib.pagesizes import letter
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                from reportlab.lib.styles import getSampleStyleSheet
                buf = io.BytesIO()
                doc = SimpleDocTemplate(buf, pagesize=letter)
                styles = getSampleStyleSheet()
                story = []
                for line in content.split('\n'):
                    line = line.strip()
                    if not line:
                        story.append(Spacer(1, 12))
                    else:
                        clean = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
                        clean = re.sub(r'\*(.+?)\*', r'\1', clean)
                        clean = re.sub(r'^#+\s*', '', clean)
                        story.append(Paragraph(clean, styles['Normal']))
                doc.build(story)
                file_bytes = buf.getvalue()
            except Exception:
                file_bytes = content.encode('utf-8')
        return file_bytes, f"{base_name}.pdf", "application/pdf"
    elif ext == "xlsx":
        try:
            from openpyxl import Workbook
            wb = Workbook()
            ws = wb.active
            # Try parsing as JSON array
            try:
                import json as _json
                data = _json.loads(content)
                if isinstance(data, list) and len(data) > 0:
                    headers = list(data[0].keys())
                    ws.append(headers)
                    for row in data:
                        ws.append([row.get(h, "") for h in headers])
            except Exception:
                # Fall back to CSV-style parsing
                for line in content.strip().split('\n'):
                    ws.append(line.split(','))
            buf = io.BytesIO()
            wb.save(buf)
            return buf.getvalue(), f"{base_name}.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        except Exception:
            return content.encode('utf-8'), f"{base_name}.csv", "text/csv"
    elif ext == "docx":
        try:
            from docx import Document
            doc = Document()
            for line in content.split('\n'):
                line = line.strip()
                if line.startswith('# '): doc.add_heading(line[2:], level=1)
                elif line.startswith('## '): doc.add_heading(line[3:], level=2)
                elif line.startswith('- '): doc.add_paragraph(line[2:], style='List Bullet')
                elif line: doc.add_paragraph(line)
            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue(), f"{base_name}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        except Exception:
            return content.encode('utf-8'), f"{base_name}.txt", "text/plain"
    elif ext == "csv":
        return content.encode('utf-8'), f"{base_name}.csv", "text/csv"
    else:
        return content.encode('utf-8'), f"{base_name}.txt", "text/plain"


# ─── PRESENTATION ─────────────────────────────────────────────────────────

class PresentationRequest(BaseModel):
    description: str
    slides: Optional[int] = 5


@router.post("/create-presentation")
async def create_presentation(req: PresentationRequest, authorization: str = Header(...), db: AsyncSession = Depends(get_db)):
    payload = _verify_auth(authorization)

    system = f"""Create a {req.slides}-slide presentation. Return JSON array where each item has "title" and "bullets" (array of strings). Example:
[{{"title":"Introduction","bullets":["Point 1","Point 2"]}},{{"title":"Details","bullets":["Info 1","Info 2"]}}]
Return ONLY valid JSON."""

    content = await _ask_claude(req.description, system)

    try:
        slides_data = json.loads(content)
    except:
        slides_data = [{"title": "Presentation", "bullets": [content]}]

    # Create simple HTML presentation (escape all AI-generated content)
    import html as _html
    html_slides = []
    for i, slide in enumerate(slides_data):
        bullets_html = "".join(f"<li>{_html.escape(str(b))}</li>" for b in slide.get("bullets", []))
        title = _html.escape(str(slide.get('title', f'Slide {i+1}')))
        html_slides.append(f"""
        <div style="page-break-after: always; padding: 60px; font-family: -apple-system, sans-serif;">
            <h1 style="font-size: 36px; color: #1a1a1a; margin-bottom: 24px;">{title}</h1>
            <ul style="font-size: 20px; line-height: 2; color: #444;">{bullets_html}</ul>
            <p style="position: absolute; bottom: 30px; right: 40px; color: #999; font-size: 14px;">{i+1} / {len(slides_data)}</p>
        </div>""")

    html = f"<html><body>{''.join(html_slides)}</body></html>"
    file_bytes = html.encode('utf-8')
    file_id = str(uuid.uuid4())
    filename = f"presentation_{file_id[:8]}.html"

    await _store_file(file_id, filename, "text/html", file_bytes, owner_email=payload.get("email", ""))

    await _audit_log_lazy()(db, payload.get("email", ""), "tool_create_presentation", "Created slide deck")
    await db.commit()
    return {"file_id": file_id, "filename": filename, "download_url": f"/api/ghost/tools/download/{file_id}", "slides": len(slides_data)}


# ─── WEB SEARCH ───────────────────────────────────────────────────────────

SERPER_API_KEY = os.getenv("SERPER_API_KEY", "")


class WebSearchRequest(BaseModel):
    query: str
    count: Optional[int] = 8


@router.post("/web-search")
async def web_search(req: WebSearchRequest, authorization: str = Header(...), db: AsyncSession = Depends(get_db)):
    payload = _verify_auth(authorization)
    results = []

    if SERPER_API_KEY:
        # ── Serper.dev — Google Search Results ──
        try:
            async with httpx.AsyncClient(timeout=15) as client:
                res = await client.post(
                    "https://google.serper.dev/search",
                    json={"q": req.query, "num": min(req.count or 8, 20)},
                    headers={"X-API-KEY": SERPER_API_KEY, "Content-Type": "application/json"},
                )
                if res.status_code == 200:
                    data = res.json()

                    # Knowledge graph (instant answer)
                    kg = data.get("knowledgeGraph", {})
                    if kg and kg.get("description"):
                        results.append({
                            "title": kg.get("title", "Knowledge Panel"),
                            "snippet": kg.get("description", ""),
                            "url": kg.get("descriptionLink") or kg.get("website", ""),
                            "type": "infobox",
                        })

                    # Organic results
                    for item in data.get("organic", [])[:8]:
                        results.append({
                            "title": item.get("title", ""),
                            "snippet": item.get("snippet", ""),
                            "url": item.get("link", ""),
                            "position": item.get("position", 0),
                        })

                    # People also ask
                    for item in data.get("peopleAlsoAsk", [])[:2]:
                        if item.get("snippet"):
                            results.append({
                                "title": item.get("question", ""),
                                "snippet": item.get("snippet", ""),
                                "url": item.get("link", ""),
                                "type": "related",
                            })

                    logger.info("[web-search] Serper returned %d results for '%s'", len(results), req.query[:60])
                else:
                    logger.warning("[web-search] Serper API returned %d: %s", res.status_code, res.text[:200])
        except Exception as e:
            logger.error("[web-search] Serper search failed: %s", e)

    if not results:
        # Fallback: DuckDuckGo instant answers (no key needed)
        try:
            from urllib.parse import quote
            async with httpx.AsyncClient(timeout=10) as client:
                res = await client.get(f"https://api.duckduckgo.com/?q={quote(req.query)}&format=json&no_html=1")
                data = res.json()
            if data.get("Abstract"):
                results.append({"title": data.get("Heading", "Answer"), "snippet": data["Abstract"], "url": data.get("AbstractURL", "")})
            for topic in data.get("RelatedTopics", [])[:5]:
                if isinstance(topic, dict) and topic.get("Text"):
                    results.append({"title": topic.get("Text", "")[:80], "snippet": topic.get("Text", ""), "url": topic.get("FirstURL", "")})
        except Exception:
            pass

    if not results:
        # Last resort: AI answer from training data
        answer = await _ask_claude(f"Answer this search query concisely with up-to-date information: {req.query}")
        results.append({"title": "AI Answer", "snippet": answer, "url": ""})

    await _audit_log_lazy()(db, payload.get("email", ""), "tool_web_search", f"Query: {req.query[:100]}")
    await db.commit()
    return {"query": req.query, "results": results}


# ─── URL READER ───────────────────────────────────────────────────────────

class ReadURLRequest(BaseModel):
    url: str
    question: Optional[str] = "Summarize this page."


@router.post("/read-url")
async def read_url(req: ReadURLRequest, authorization: str = Header(...), db: AsyncSession = Depends(get_db)):
    payload = _verify_auth(authorization)
    # SSRF protection — block internal/private URLs
    from routes.ghost_tools_v2 import _validate_external_url
    _validate_external_url(req.url)

    async with httpx.AsyncClient(timeout=15, follow_redirects=True) as client:
        res = await client.get(req.url, headers={"User-Agent": "GoFarther-AI/1.0"})
        if res.status_code != 200:
            raise HTTPException(400, f"Could not fetch URL (status {res.status_code})")
        html = res.text[:50000]  # Limit to 50KB

    # Strip HTML tags (basic)
    import re
    text = re.sub(r'<script[^>]*>.*?</script>', '', html, flags=re.DOTALL)
    text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL)
    text = re.sub(r'<[^>]+>', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()[:10000]

    # Ask AI to answer based on content
    summary = await _ask_claude(f"Based on this webpage content, {req.question}\n\nContent:\n{text}")

    await _audit_log_lazy()(db, payload.get("email", ""), "tool_read_url", f"URL: {req.url[:100]}")
    await db.commit()
    return {"url": req.url, "summary": summary}


# ─── PDF READER ───────────────────────────────────────────────────────────

class ReadPDFRequest(BaseModel):
    pdf_base64: str
    question: Optional[str] = "Summarize this document."


@router.post("/read-pdf")
async def read_pdf(req: ReadPDFRequest, authorization: str = Header(...), db: AsyncSession = Depends(get_db)):
    payload = _verify_auth(authorization)

    try:
        import fitz  # PyMuPDF
        pdf_bytes = base64.b64decode(req.pdf_base64)
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        text = text[:10000]
    except ImportError:
        raise HTTPException(500, "PDF reader not available on server")
    except Exception as e:
        raise HTTPException(400, f"Could not read PDF: {str(e)}")

    summary = await _ask_claude(f"{req.question}\n\nDocument content:\n{text}")
    await _audit_log_lazy()(db, payload.get("email", ""), "tool_read_pdf", "PDF read")
    await db.commit()
    return {"summary": summary, "pages": len(doc) if doc else 0}


# ─── STOCK REPORT ─────────────────────────────────────────────────────────

class StockRequest(BaseModel):
    symbol: str


@router.post("/stock-report")
async def stock_report(req: StockRequest, authorization: str = Header(...), db: AsyncSession = Depends(get_db)):
    payload = _verify_auth(authorization)

    symbol = req.symbol.upper()
    async with httpx.AsyncClient(timeout=15) as client:
        res = await client.get(f"https://query1.finance.yahoo.com/v8/finance/chart/{symbol}?interval=1d&range=1mo")
        data = res.json()

    meta = data.get("chart", {}).get("result", [{}])[0].get("meta", {})
    price = meta.get("regularMarketPrice", "N/A")
    prev_close = meta.get("previousClose", "N/A")
    currency = meta.get("currency", "USD")

    report = await _ask_claude(f"Give a brief stock analysis for {symbol}. Current price: ${price} {currency}. Previous close: ${prev_close}. Include a brief outlook.")

    await _audit_log_lazy()(db, payload.get("email", ""), "tool_stock_report", f"Symbol: {symbol}")
    await db.commit()
    return {"symbol": symbol, "price": price, "currency": currency, "previous_close": prev_close, "analysis": report}


# ─── WEATHER REPORT ───────────────────────────────────────────────────────

class WeatherRequest(BaseModel):
    location: str


@router.post("/weather-report")
async def weather_report(req: WeatherRequest, authorization: str = Header(...), db: AsyncSession = Depends(get_db)):
    payload = _verify_auth(authorization)

    async with httpx.AsyncClient(timeout=10) as client:
        from urllib.parse import quote
        res = await client.get(f"https://wttr.in/{quote(req.location, safe='')}?format=j1")
        data = res.json()

    current = data.get("current_condition", [{}])[0]
    forecast = data.get("weather", [])[:3]

    weather_info = {
        "location": req.location,
        "temperature": f"{current.get('temp_F', '?')}°F / {current.get('temp_C', '?')}°C",
        "condition": current.get("weatherDesc", [{}])[0].get("value", "Unknown"),
        "humidity": f"{current.get('humidity', '?')}%",
        "wind": f"{current.get('windspeedMiles', '?')} mph",
        "forecast": [{"date": d.get("date", ""), "high": f"{d.get('maxtempF', '?')}°F", "low": f"{d.get('mintempF', '?')}°F"} for d in forecast],
    }

    await _audit_log_lazy()(db, payload.get("email", ""), "tool_weather_report", f"Location: {req.location}")
    await db.commit()
    return weather_info


# ─── NEWS ─────────────────────────────────────────────────────────────────

class NewsRequest(BaseModel):
    topic: str


@router.post("/news")
async def news(req: NewsRequest, authorization: str = Header(...), db: AsyncSession = Depends(get_db)):
    payload = _verify_auth(authorization)

    # Use DuckDuckGo news
    async with httpx.AsyncClient(timeout=15) as client:
        from urllib.parse import quote
        res = await client.get(f"https://api.duckduckgo.com/?q={quote(req.topic)}&format=json&no_html=1")
        data = res.json()

    # Get AI to provide news summary based on its knowledge
    summary = await _ask_claude(f"Give me the latest news and developments about: {req.topic}. Be specific with recent events, dates, and key details. Format as bullet points.")

    await _audit_log_lazy()(db, payload.get("email", ""), "tool_news", f"Topic: {req.topic[:100]}")
    await db.commit()
    return {"topic": req.topic, "summary": summary}


# ─── TRANSLATE DOCUMENT ───────────────────────────────────────────────────

class TranslateRequest(BaseModel):
    text: str
    target_language: str  # e.g. "Spanish", "French"


@router.post("/translate-doc")
async def translate_doc(req: TranslateRequest, authorization: str = Header(...), db: AsyncSession = Depends(get_db)):
    payload = _verify_auth(authorization)

    translation = await _ask_claude(
        f"Translate the following text to {req.target_language}. Return ONLY the translation, nothing else.\n\n{req.text}",
        system=f"You are a professional translator. Translate accurately to {req.target_language}."
    )

    await _audit_log_lazy()(db, payload.get("email", ""), "tool_translate_doc", f"To: {req.target_language}")
    await db.commit()
    return {"original_language": "auto-detected", "target_language": req.target_language, "translation": translation}


# ─── CODE INTERPRETER ─────────────────────────────────────────────────────

class RunCodeRequest(BaseModel):
    description: str  # What the user wants to compute/create


@router.post("/run-code")
async def run_code(req: RunCodeRequest, authorization: str = Header(...), db: AsyncSession = Depends(get_db)):
    payload = _verify_auth(authorization)

    # Ask AI to write Python code
    code = await _ask_claude(
        f"Write Python code to: {req.description}\n\nReturn ONLY the Python code, no explanations. Use print() for output. Do not use any dangerous operations (no file deletion, no network requests, no system commands).",
        system="You are a Python code generator. Write safe, clean Python code. Only use standard library modules."
    )

    # Clean the code (remove markdown fences)
    code = code.strip()
    if code.startswith("```python"):
        code = code[9:]
    if code.startswith("```"):
        code = code[3:]
    if code.endswith("```"):
        code = code[:-3]
    code = code.strip()

    # AST-based security check — whitelist safe modules, block dangerous builtins
    import ast as _ast
    _ALLOWED_MODULES = {"math", "statistics", "json", "datetime", "re", "collections",
                        "itertools", "functools", "decimal", "fractions", "random",
                        "string", "textwrap", "csv", "operator", "numbers"}
    _BLOCKED_BUILTINS = {"__import__", "exec", "eval", "compile", "open", "getattr",
                         "setattr", "delattr", "globals", "locals", "vars", "dir",
                         "type", "breakpoint", "input", "memoryview", "help"}

    try:
        tree = _ast.parse(code)
    except SyntaxError as se:
        return {"code": code, "output": f"Syntax error in generated code: {se}"}

    for node in _ast.walk(tree):
        if isinstance(node, _ast.Import):
            for alias in node.names:
                mod = alias.name.split(".")[0]
                if mod not in _ALLOWED_MODULES:
                    return {"code": code, "output": f"Blocked: import of '{mod}' is not allowed for security reasons"}
        elif isinstance(node, _ast.ImportFrom):
            if node.module:
                mod = node.module.split(".")[0]
                if mod not in _ALLOWED_MODULES:
                    return {"code": code, "output": f"Blocked: import from '{mod}' is not allowed for security reasons"}
        elif isinstance(node, _ast.Call):
            func = node.func
            if isinstance(func, _ast.Name) and func.id in _BLOCKED_BUILTINS:
                return {"code": code, "output": f"Blocked: '{func.id}()' is not allowed for security reasons"}
            elif isinstance(func, _ast.Attribute) and func.attr in _BLOCKED_BUILTINS:
                return {"code": code, "output": f"Blocked: '.{func.attr}()' is not allowed for security reasons"}

    # Execute in a restricted subprocess — minimal env, no HOME
    import subprocess
    try:
        result = subprocess.run(
            ["python3", "-c", code],
            capture_output=True, text=True, timeout=10,
            env={"PATH": "/usr/bin:/usr/local/bin"}
        )
        output = result.stdout or result.stderr or "No output"
    except subprocess.TimeoutExpired:
        output = "Code execution timed out (10s limit)"
    except Exception as e:
        logger.error("Code execution error: %s", e)
        output = "Code execution failed"

    await _audit_log_lazy()(db, payload.get("email", ""), "tool_run_code", "Python code execution")
    await db.commit()
    return {"code": code, "output": output.strip()}


# ─── DATA ANALYZER ────────────────────────────────────────────────────────

class AnalyzeDataRequest(BaseModel):
    csv_data: str  # Raw CSV content
    question: Optional[str] = "Analyze this data and provide key insights."


@router.post("/analyze-data")
async def analyze_data(req: AnalyzeDataRequest, authorization: str = Header(...), db: AsyncSession = Depends(get_db)):
    payload = _verify_auth(authorization)

    # Limit data size
    csv_preview = req.csv_data[:5000]

    analysis = await _ask_claude(
        f"{req.question}\n\nCSV Data:\n{csv_preview}",
        system="You are a data analyst. Analyze the CSV data provided. Give clear insights, trends, and statistics. Format with bullet points and headers."
    )

    await _audit_log_lazy()(db, payload.get("email", ""), "tool_analyze_data", "Data analysis")
    await db.commit()
    return {"analysis": analysis, "rows_analyzed": req.csv_data.count('\n')}
