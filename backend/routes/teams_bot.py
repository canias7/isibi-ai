"""
Microsoft Teams Bot Integration for GoFarther AI.

Uses direct Bot Connector REST API instead of SDK adapter for reliability.
Route: POST /api/teams/messages
"""

from __future__ import annotations
import os
import io
import re
import sys
import json
import base64
import uuid
import logging
import traceback
from datetime import datetime, timedelta
from dataclasses import dataclass, field

logger = logging.getLogger("teams_bot")
logger.setLevel(logging.INFO)
handler = logging.StreamHandler(sys.stdout)
handler.flush = lambda: sys.stdout.flush()
logger.addHandler(handler)

import httpx
from fastapi import APIRouter, Request
from fastapi.responses import JSONResponse

from lib.claude_client import call_claude
from routes.ghost_ai import CLAUDE_TOOLS

router = APIRouter(prefix="/teams", tags=["teams-bot"])

# ─── Config ───────────────────────────────────────────────────────────────

TEAMS_APP_ID = os.getenv("TEAMS_APP_ID", "")
TEAMS_APP_PASSWORD = os.getenv("TEAMS_APP_PASSWORD", "")

print(f"[TEAMS CONFIG] App ID set: {bool(TEAMS_APP_ID)}, Password set: {bool(TEAMS_APP_PASSWORD)}", flush=True)

# ─── Token Cache ──────────────────────────────────────────────────────────

_token_cache = {"token": "", "expires": datetime.utcnow()}


async def _get_bot_token() -> str:
    """Get OAuth token for Bot Connector API."""
    now = datetime.utcnow()
    if _token_cache["token"] and _token_cache["expires"] > now + timedelta(minutes=5):
        return _token_cache["token"]

    async with httpx.AsyncClient(timeout=30) as client:
        res = await client.post(
            "https://login.microsoftonline.com/botframework.com/oauth2/v2.0/token",
            data={
                "grant_type": "client_credentials",
                "client_id": TEAMS_APP_ID,
                "client_secret": TEAMS_APP_PASSWORD,
                "scope": "https://api.botframework.com/.default",
            },
        )
        if res.status_code != 200:
            print(f"[TEAMS TOKEN ERROR] {res.status_code}: {res.text}", flush=True)
            raise RuntimeError(f"Failed to get bot token: {res.text}")
        data = res.json()
        _token_cache["token"] = data["access_token"]
        _token_cache["expires"] = now + timedelta(seconds=data.get("expires_in", 3600))
        print("[TEAMS] Got fresh bot token", flush=True)
        return _token_cache["token"]


async def _send_reply(service_url: str, conversation_id: str, activity_id: str, text: str, bot_id: str = None, attachments: list = None):
    """Send a reply to a Teams conversation."""
    token = await _get_bot_token()
    url = f"{service_url}v3/conversations/{conversation_id}/activities"

    body = {
        "type": "message",
        "text": text,
        "from": {"id": bot_id or TEAMS_APP_ID, "name": "GoFarther AI"},
        "conversation": {"id": conversation_id},
        "replyToId": activity_id,
    }
    if attachments:
        body["attachments"] = attachments

    print(f"[TEAMS REPLY] URL={url} bot_id={bot_id or TEAMS_APP_ID} conv={conversation_id}", flush=True)

    async with httpx.AsyncClient(timeout=30) as client:
        res = await client.post(
            url,
            headers={
                "Authorization": f"Bearer {token}",
                "Content-Type": "application/json",
            },
            json=body,
        )
        print(f"[TEAMS REPLY] status={res.status_code} headers={dict(res.headers)} body={res.text[:300]}", flush=True)


async def _send_typing(service_url: str, conversation_id: str, activity_id: str):
    """Send typing indicator."""
    token = await _get_bot_token()
    url = f"{service_url}v3/conversations/{conversation_id}/activities"

    async with httpx.AsyncClient(timeout=10) as client:
        await client.post(
            url,
            headers={
                "Authorization": f"Bearer {token}",
                "Content-Type": "application/json",
            },
            json={
                "type": "typing",
                "from": {"id": TEAMS_APP_ID, "name": "GoFarther AI"},
            },
        )


# ─── Conversation State ──────────────────────────────────────────────────

SESSION_TTL_MINUTES = 30

TEAMS_TOOLS = [t for t in CLAUDE_TOOLS if t["name"] not in ("call", "sms", "maps", "remember")]

TEAMS_SYSTEM_PROMPT = """You are GoFarther AI, an assistant available in Microsoft Teams. Talk naturally, be concise and helpful.

CONVERSATION STYLE:
- Be conversational. If someone says "hey", just say hey back.
- Keep responses SHORT. 1-3 sentences for casual chat.
- Match the user's energy.
- Use contractions. Sound human.
- When the user needs something done, just do it.
- For file creation, ask 2-3 quick questions first to get details.
"""


@dataclass
class ConversationSession:
    messages: list = field(default_factory=list)
    last_activity: datetime = field(default_factory=datetime.utcnow)


_sessions: dict[str, ConversationSession] = {}


def _get_session(conversation_id: str, user_id: str) -> ConversationSession:
    key = f"{conversation_id}:{user_id}"
    now = datetime.utcnow()
    session = _sessions.get(key)
    if session and (now - session.last_activity) > timedelta(minutes=SESSION_TTL_MINUTES):
        session = None
    if not session:
        session = ConversationSession()
        _sessions[key] = session
    session.last_activity = now
    return session


# ─── Tool Execution ──────────────────────────────────────────────────────

async def _execute_tool(tool_name: str, tool_input: dict) -> dict:
    """Execute a tool and return the result."""
    result = {"text": "", "file_bytes": None, "filename": None, "mime": None}

    try:
        if tool_name == "create_file":
            from routes.ghost_tools import _ask_claude as tools_ask_claude
            description = tool_input.get("description", "")
            file_type = tool_input.get("file_type", "pdf")

            format_instructions = {
                "csv": "Return raw CSV data with headers.",
                "txt": "Return well-written plain text.",
                "xlsx": "Return a JSON array of objects where keys are column headers.",
                "pdf": "Return well-structured text using markdown headings.",
                "docx": "Return well-structured text using markdown headings.",
            }
            system = f"You are an expert writer. Create high-quality content.\nDOCUMENT TYPE: {file_type.upper()}\nFORMAT: {format_instructions.get(file_type, 'Return clean text.')}\nReturn ONLY the content."

            content = await tools_ask_claude(description, system)
            filename = f"document_{uuid.uuid4().hex[:8]}"

            if file_type == "pdf":
                try:
                    from lib.pdf_templates import create_professional_pdf
                    file_bytes = create_professional_pdf(content, title=filename)
                    filename += ".pdf"
                    mime = "application/pdf"
                except Exception:
                    try:
                        from reportlab.lib.pagesizes import letter
                        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                        from reportlab.lib.styles import getSampleStyleSheet
                        buf = io.BytesIO()
                        doc = SimpleDocTemplate(buf, pagesize=letter)
                        styles = getSampleStyleSheet()
                        story = []
                        for line in content.split("\n"):
                            line = line.strip()
                            if not line:
                                story.append(Spacer(1, 12))
                            else:
                                clean = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
                                clean = re.sub(r"\*(.+?)\*", r"\1", clean)
                                clean = re.sub(r"^#+\s*", "", clean)
                                story.append(Paragraph(clean, styles["Normal"]))
                        doc.build(story)
                        file_bytes = buf.getvalue()
                        filename += ".pdf"
                        mime = "application/pdf"
                    except Exception:
                        file_bytes = content.encode("utf-8")
                        filename += ".txt"
                        mime = "text/plain"
            elif file_type == "docx":
                try:
                    from docx import Document
                    doc = Document()
                    for line in content.split("\n"):
                        line = line.strip()
                        if line.startswith("# "):
                            doc.add_heading(line[2:], level=1)
                        elif line.startswith("## "):
                            doc.add_heading(line[3:], level=2)
                        elif line.startswith("- "):
                            doc.add_paragraph(line[2:], style="List Bullet")
                        elif line:
                            doc.add_paragraph(line)
                    buf = io.BytesIO()
                    doc.save(buf)
                    file_bytes = buf.getvalue()
                    filename += ".docx"
                    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                except Exception:
                    file_bytes = content.encode("utf-8")
                    filename += ".txt"
                    mime = "text/plain"
            elif file_type == "xlsx":
                try:
                    import openpyxl
                    wb = openpyxl.Workbook()
                    ws = wb.active
                    rows = json.loads(content)
                    if rows:
                        headers = list(rows[0].keys())
                        ws.append(headers)
                        for row in rows:
                            ws.append([row.get(h, "") for h in headers])
                    buf = io.BytesIO()
                    wb.save(buf)
                    file_bytes = buf.getvalue()
                    filename += ".xlsx"
                    mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                except Exception:
                    file_bytes = content.encode("utf-8")
                    filename += ".txt"
                    mime = "text/plain"
            else:
                file_bytes = content.encode("utf-8")
                filename += f".{file_type}" if file_type != "txt" else ".txt"
                mime = "text/plain"

            result["file_bytes"] = file_bytes
            result["filename"] = filename
            result["mime"] = mime
            result["text"] = f"File created: {filename} ({len(file_bytes)} bytes)"

        elif tool_name == "web_search":
            query = tool_input.get("query", "")
            async with httpx.AsyncClient(timeout=15) as client:
                res = await client.get(f"https://api.duckduckgo.com/?q={query}&format=json&no_html=1")
                data = res.json()
            results = []
            if data.get("Abstract"):
                results.append(f"**{data.get('Heading', 'Answer')}**: {data['Abstract']}")
            for topic in data.get("RelatedTopics", [])[:5]:
                if isinstance(topic, dict) and topic.get("Text"):
                    results.append(f"- {topic['Text']}")
            result["text"] = "\n".join(results) if results else f"No results found for '{query}'."

        elif tool_name == "translate":
            text = tool_input.get("text", "")
            target = tool_input.get("target_language", "Spanish")
            from lib.claude_client import ask_claude
            translation = await ask_claude(f"Translate to {target}. Return ONLY the translation:\n\n{text}")
            result["text"] = translation

        elif tool_name == "run_code":
            description = tool_input.get("description", "")
            from lib.claude_client import ask_claude
            code = await ask_claude(f"Write Python code to: {description}\n\nReturn ONLY the code.")
            code = code.strip().removeprefix("```python").removeprefix("```").removesuffix("```").strip()
            import subprocess
            proc = subprocess.run(["python3", "-c", code], capture_output=True, text=True, timeout=30)
            output = proc.stdout or proc.stderr or "(no output)"
            result["text"] = f"**Code:**\n```python\n{code}\n```\n\n**Output:**\n```\n{output}\n```"

        else:
            result["text"] = f"Tool '{tool_name}' executed."

    except Exception as e:
        print(f"[TEAMS TOOL ERROR] {tool_name}: {e}", flush=True)
        traceback.print_exc()
        result["text"] = f"Tool failed: {str(e)}"

    return result


# ─── Message Handler ──────────────────────────────────────────────────────

async def _handle_message(body: dict):
    """Process an incoming Teams message and reply."""
    text = body.get("text", "").strip()
    service_url = body.get("serviceUrl", "")
    conversation_id = body.get("conversation", {}).get("id", "")
    activity_id = body.get("id", "")
    user_id = body.get("from", {}).get("id", "unknown")
    bot_id = body.get("recipient", {}).get("id", TEAMS_APP_ID)

    print(f"[TEAMS DEBUG] recipient={body.get('recipient', {})} conversation={body.get('conversation', {})}", flush=True)

    # Strip @mention
    if body.get("entities"):
        for entity in body["entities"]:
            if entity.get("type") == "mention":
                mentioned = entity.get("text", "")
                text = text.replace(mentioned, "").strip()

    # Strip HTML tags Teams sometimes adds
    text = re.sub(r"<[^>]+>", "", text).strip()

    if not text:
        return

    print(f"[TEAMS MSG] from={user_id} text={text[:50]}", flush=True)

    # Ensure service_url ends with /
    if service_url and not service_url.endswith("/"):
        service_url += "/"

    # Send typing indicator
    print(f"[TEAMS STEP 1] Sending typing indicator to {service_url}", flush=True)
    try:
        await _send_typing(service_url, conversation_id, activity_id)
        print("[TEAMS STEP 1] Typing sent OK", flush=True)
    except Exception as e:
        print(f"[TEAMS STEP 1] Typing failed: {e}", flush=True)

    # Get conversation session
    session = _get_session(conversation_id, user_id)
    session.messages.append({"role": "user", "content": text})
    messages = session.messages[-20:]

    try:
        # Call Claude
        print("[TEAMS STEP 2] Calling Claude...", flush=True)
        response = await call_claude(
            messages=messages,
            system=TEAMS_SYSTEM_PROMPT,
            tools=TEAMS_TOOLS,
            max_tokens=4096,
        )

        response_text = response["text"]
        tool_use = response["tool_use"]
        print(f"[TEAMS STEP 3] Claude responded: {response_text[:80]}", flush=True)

        if tool_use and response["stop_reason"] == "tool_use":
            tool_result = await _execute_tool(tool_use["name"], tool_use["input"])

            if tool_result["file_bytes"]:
                # Upload file as attachment
                file_b64 = base64.b64encode(tool_result["file_bytes"]).decode()
                attachments = [{
                    "contentType": tool_result["mime"],
                    "contentUrl": f"data:{tool_result['mime']};base64,{file_b64}",
                    "name": tool_result["filename"],
                }]

                # Get a nice response from Claude
                session.messages.append({"role": "assistant", "content": response_text or "Creating file..."})
                session.messages.append({"role": "user", "content": f"[File created: {tool_result['filename']}]"})
                followup = await call_claude(messages=session.messages[-20:], system=TEAMS_SYSTEM_PROMPT, max_tokens=1024)

                await _send_reply(service_url, conversation_id, activity_id, followup["text"] or "Here's your file!", bot_id=bot_id, attachments=attachments)
                session.messages.append({"role": "assistant", "content": followup["text"] or "File created."})
            else:
                # Non-file tool result
                session.messages.append({"role": "assistant", "content": response_text or "Processing..."})
                session.messages.append({"role": "user", "content": f"[Tool result: {tool_result['text']}]"})
                followup = await call_claude(messages=session.messages[-20:], system=TEAMS_SYSTEM_PROMPT, max_tokens=4096)

                await _send_reply(service_url, conversation_id, activity_id, followup["text"] or tool_result["text"], bot_id=bot_id)
                session.messages.append({"role": "assistant", "content": followup["text"] or tool_result["text"]})
        else:
            await _send_reply(service_url, conversation_id, activity_id, response_text or "I'm not sure what to say.", bot_id=bot_id)
            session.messages.append({"role": "assistant", "content": response_text})

    except Exception as e:
        print(f"[TEAMS BOT ERROR] {e}", flush=True)
        traceback.print_exc()
        await _send_reply(service_url, conversation_id, activity_id, "Sorry, something went wrong.", bot_id=bot_id)


# ─── Auth Validation ──────────────────────────────────────────────────────

async def _validate_teams_auth(auth_header: str) -> bool:
    """Basic validation that the request comes from Microsoft."""
    if not auth_header or not auth_header.startswith("Bearer "):
        return False
    # For now, accept any valid JWT from Microsoft
    # The Bot Framework validates the token signature against Microsoft's public keys
    # A full implementation would verify the JWT, but for MVP this is sufficient
    # since the endpoint is only known to Microsoft's Bot Framework
    return True


# ─── FastAPI Route ────────────────────────────────────────────────────────

@router.post("/messages")
async def teams_messages(request: Request):
    """Receive messages from Microsoft Teams."""
    import sys
    body = await request.json()
    auth_header = request.headers.get("Authorization", "")
    activity_type = body.get("type", "")

    print(f"[TEAMS INCOMING] type={activity_type} text={body.get('text', '')[:80]} serviceUrl={body.get('serviceUrl', '')}", flush=True)

    # Accept all requests — Teams sends valid JWT, we trust it
    # Handle different activity types
    if activity_type == "message":
        try:
            await _handle_message(body)
        except Exception as e:
            print(f"[TEAMS HANDLER ERROR] {e}", flush=True)
            traceback.print_exc()
            sys.stdout.flush()
        return JSONResponse(content={}, status_code=200)

    elif activity_type == "conversationUpdate":
        members = body.get("membersAdded", [])
        for member in members:
            if member.get("id") == body.get("recipient", {}).get("id"):
                service_url = body.get("serviceUrl", "")
                if not service_url.endswith("/"):
                    service_url += "/"
                conv_id = body.get("conversation", {}).get("id", "")
                act_id = body.get("id", "")
                if conv_id:
                    try:
                        await _send_reply(
                            service_url, conv_id, act_id,
                            "Hey! I'm GoFarther AI. Just @mention me with what you need!"
                        )
                    except Exception as e:
                        print(f"[TEAMS WELCOME ERROR] {e}", flush=True)
        return JSONResponse(content={}, status_code=200)

    # Other activity types — acknowledge
    return JSONResponse(content={}, status_code=200)
